import json, re, time
from datetime import datetime
from tqdm import tqdm
from docx import Document
from concurrent.futures import ThreadPoolExecutor

from deepseek_client import DeepSeekClient
from document_utils import setup_document_style, add_answer_section
from conf.config import DEEPSEEK_API_KEY

class EnhancedInterviewGenerator:
    def __init__(self, tech_direction="JAVA"):
        self.tech_direction = tech_direction
        self.api_key = DEEPSEEK_API_KEY
        self.base_url = "https://api.deepseek.com"
        self.doc = Document()
        # 存放所有题目的答案与解析信息，后续用于生成参考答案与解析部分
        self.answer_sheet = []  
        # 用于记录所有题目的编号（全局递增）
        self.question_count = 0  
        # 存放各题型生成的题目数据，结构：{ "单选题": [q1, q2, ...], "多选题": [...], ... }
        self.generated_questions = {}  
        self.deepseek_client = DeepSeekClient(api_key=self.api_key, base_url=self.base_url, tech_direction=self.tech_direction)

    def _get_filename(self):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        return f"{self.tech_direction}_Interview_{timestamp}.docx"
    
    def _generate_batch_questions(self, question_type, num=15):
        """
        调用 API 生成一批题目，并对生成的题目进行格式校验（多选题单个题目错误则剔除）
        若遇到部分题目错误，则仅剔除错误题目，返回格式正确的题目列表。
        """
        # 定义各题型要求
        type_requirements = {
            "单选题": "每个问题必须包含四个中文选项，用大写字母A. B. C. D. 标记，答案只能有一个正确选项",
            "多选题": "每个问题必须包含5-7个中文选项，用大写字母A. B. C. D. E. F...标记，答案应有2-7个正确选项。问题内容必须明确列出所有选项，每个选项用大写字母开头（例如：A. 选项内容）",
            "填空题": "问题中用___________表示空白，答案填具体内容",
            "判断题": "答案只能是'正确'或'错误'",
            "问答题": "问题必须是不含任何选项的开放式技术问题，答案为简明扼要的文字描述"
        }
        # 定义示例文本
        question_examples = {
            "单选题": '"问题内容（示例：\'Java中的final关键字作用？\\nA. 继承\\nB. 重写\\nC. 常量\\nD. 多态\'）"',
            "多选题": '"问题内容（示例：\'哪些是Java集合接口？\\nA. List\\nB. Set\\nC. Map\\nD. Array\'）"',
            "问答题": '问题内容（示例："请详细解释C#中垃圾回收机制及其优化方法"）'
        }
        
        if question_type == "问答题":
            prompt = f"""请生成{num}道{self.tech_direction}高级开发工程师面试{question_type}，
                要求：
                1. {type_requirements[question_type]}。
                2. 所有问题和解析必须使用中文（专有名称和术语除外）。
                3. 使用严格的JSON格式，仅返回JSON，不包含任何其他文本，格式如下：
                {{
                    "questions": [
                        {{
                            "question": "请详细解释{self.tech_direction}中某个关键技术原理",
                            "short_answer": "最简答案",
                            "detailed_analysis": "300字技术解析",
                            "keywords": ["关键词"],
                            "difficulty": "1-5"
                        }}
                    ]
                }}
                4. 请确保生成的题目各不相同，不要包含任何选项。
            """
        else:
            prompt = f"""请生成{num}道{self.tech_direction}高级开发工程师面试{question_type}，要求：
                1. {type_requirements[question_type]}
                2. 所有问题、选项和解析必须使用中文（专有名称和术语除外）
                3. {"多选题答案字母按升序排列（如ABE）" if question_type == "多选题" else ""}
                4. 使用严格JSON格式：
                {{
                    "questions": [
                        {{
                            "question": {question_examples.get(question_type, '"问题内容"')},
                            "short_answer": "最简答案",
                            "detailed_analysis": "300字技术解析",
                            "keywords": ["关键词"],
                            "difficulty": "1-5"
                        }}
                    ]
                }}
                5. 不要包含多余文本
                6. 请确保生成的题目各不相同
            """
        
        result = self.deepseek_client.call(prompt)
        if not result:
            return []
        result = re.sub(r'```json|```', '', result).strip()
        try:
            data = json.loads(result)
            questions = data.get("questions", [])[:num]
            valid_questions = []
            if question_type == "单选题":
                for q in questions:
                    valid = True
                    answer = q.get("short_answer", "")
                    if len(answer) != 1 or not answer.isalpha():
                        print(f"单选题答案格式错误: {answer}，已剔除该题")
                        valid = False
                    options = re.findall(r'\b[A-Z]\.\s', q.get("question", ""))
                    if len(options) != 4:
                        print(f"单选题选项数量不为4: {q.get('question','')}，已剔除该题")
                        valid = False
                    if valid:
                        valid_questions.append(q)
            elif question_type == "多选题":
                for q in questions:
                    valid = True
                    answer = q.get("short_answer", "")
                    if len(answer) < 2 or not answer.isalpha() or answer != ''.join(sorted(answer)):
                        print(f"多选题答案格式错误: {answer}，已剔除该题")
                        valid = False
                    if not re.search(r'\b[A-Z]\.\s', q.get("question", "")):
                        print(f"多选题选项缺失: {q.get('question','')}，已剔除该题")
                        valid = False
                    if valid:
                        valid_questions.append(q)
            elif question_type == "问答题":
                for q in questions:
                    if re.search(r'\b[A-Z]\.\s', q.get("question", "")):
                        print(f"问答题包含选项: {q.get('question','')}，已剔除该题")
                    else:
                        valid_questions.append(q)
            else:
                valid_questions = questions
            return valid_questions
        except Exception as e:
            print(f"解析失败: {str(e)}")
            print("原始内容:", result)
            return []
    
    def _write_question_sections(self, generated_questions):
        """
        根据生成的题目数据写入文档，按照题型分组
        """
        for q_type in generated_questions:
            questions = generated_questions[q_type]
            self.doc.add_heading(f"{q_type}（共{len(questions)}题）", level=2)
            for q in questions:
                if "number" in q:
                    self.doc.add_paragraph(
                        f"{q['number']}. {q['question']}（难度：{q['difficulty']}/5）",
                        style='Normal'
                    )
                else:
                    self.doc.add_paragraph(q.get('question', '无题'), style='Normal')
            self.doc.add_page_break()
    
    def _add_knowledge_points_summary(self):
        """
        根据所有生成的题目（包括判断题），调用 API 分批生成一份非常详细的知识点总结，
        要求对每一道题分别说明其涉及的知识点，详细说明包括：
           1. 详细原理和运行机制（要求内容非常详细）
           2. 实际应用场景及具体示例
           3. 使用时的注意事项和防范措施
        每一道题的说明请按照如下模板输出，并在每道题后输出“====”作为分隔符：
           【知识点名称】：
           【原理】：
           【实际应用】：
           【注意事项】：
        如果所有题目字数过长，请分批生成，确保每道题都有详细的总结。
        """
        import re
        # 汇总所有题目
        all_questions = []
        for q_type, questions in self.generated_questions.items():
            all_questions.extend(questions)
        all_questions.sort(key=lambda x: x['number'])
        total_questions = len(all_questions)
        batch_size = 10  # 每批处理10道题，可根据需要调整
        responses = []
        print("开始分批生成详细知识点总结，请耐心等待...")
        knowledge_start = time.time()
        for i in range(0, total_questions, batch_size):
            batch = all_questions[i:i+batch_size]
            questions_text = "\n".join([f"{item['number']}. {item['question']}" for item in batch])
            prompt = (
                f"请对以下{self.tech_direction}高级开发面试题目中的每一道题分别说明其涉及的知识点，要求：\n"
                "1. 每一道题的说明应包括：详细原理和运行机制（要求内容非常详细）、实际应用场景及具体示例、使用时的注意事项和防范措施。\n"
                "2. 请对每一道题单独输出说明，并按照如下模板输出，每一道题的说明后请单独输出一行‘====’作为分隔符：\n"
                "【知识点名称】：\n"
                "【原理】：\n"
                "【实际应用】：\n"
                "【注意事项】：\n"
                "请确保所有题目的知识点均被覆盖，并输出为纯文本格式，避免使用 Markdown 语法。\n"
                "题目如下：\n"
                f"{questions_text}"
            )
            batch_response = self.deepseek_client.call(prompt)
            if batch_response:
                batch_response = re.sub(r'```.+?```', '', batch_response).strip()
                responses.append(batch_response)
            else:
                print("某批次知识点总结生成失败")
        knowledge_end = time.time()
        elapsed = knowledge_end - knowledge_start
        minutes = int(elapsed // 60)
        seconds = elapsed % 60
        print(f"详细知识点总结生成完成，耗时 {minutes}分钟{seconds:.2f}秒")
        return "\n".join(responses)
    
    def generate_exam_paper(self, question_types):
        overall_start = time.time()
        from document_utils import add_knowledge_summary_section_template
        setup_document_style(self.doc)
        self.doc.add_heading(f"{self.tech_direction}高级开发面试题库", level=0)
        self.doc.add_paragraph("\n考生姓名：__________\n考试时间：120分钟\n\n")
        
        # 存储各题型生成的题目数据
        self.generated_questions = {}
        
        # 先生成各类题目，但不直接写入文档
        for q_type, total in question_types:
            generated = []
            remaining = total
            seen_questions = set()
            attempts = 0
            max_attempts = 5
            pbar = tqdm(total=total, desc=f"生成 {q_type}", ncols=80)
            while remaining > 0 and attempts < max_attempts:
                batch_num = remaining if remaining < 10 else 10
                questions = self._generate_batch_questions(q_type, batch_num)
                if questions is None:
                    print(f"{q_type}生成失败，请检查API设置")
                    break
                unique_questions = []
                for q in questions:
                    if q["question"] not in seen_questions:
                        seen_questions.add(q["question"])
                        self.question_count += 1
                        q_with_num = q.copy()
                        q_with_num["number"] = self.question_count
                        q_with_num["type"] = q_type  # 添加题型信息
                        q_with_num["answer"] = q.get("short_answer", "")
                        q_with_num["analysis"] = q.get("detailed_analysis", "")
                        unique_questions.append(q_with_num)
                        self.answer_sheet.append(q_with_num)
                    else:
                        print(f"检测到重复题目，已跳过: {q['question']}")
                if len(unique_questions) == 0:
                    attempts += 1
                    print(f"当前批次{q_type}重复或格式错误较多，尝试补充次数：{attempts}")
                else:
                    attempts = 0
                generated.extend(unique_questions)
                pbar.update(len(unique_questions))
                remaining = total - len(generated)
            pbar.close()
            if remaining > 0:
                print(f"警告：{q_type}最终未生成足够题目，期望{total}题，实际获得{len(generated)}题")
            self.generated_questions[q_type] = generated
        
        # 根据题目生成知识点总结，并插入到考生信息之后
        knowledge_points = self._add_knowledge_points_summary()
        if knowledge_points:
            add_knowledge_summary_section_template(self.doc, knowledge_points)
            self.doc.add_page_break()
            print("知识点总结添加成功")
        else:
            print("未生成知识点总结")
        
        # 将各题型题目写入文档
        self._write_question_sections(self.generated_questions)
        
        overall_end = time.time()
        total_elapsed = overall_end - overall_start
        minutes = int(total_elapsed // 60)
        seconds = total_elapsed % 60
        add_answer_section(self.doc, self.answer_sheet)
        filename = self._get_filename()
        self.doc.save(filename)
        print(f"生成成功！文件已保存为 {filename}, 共{self.question_count}题")
        print(f"总耗时: {minutes}分钟{seconds:.2f}秒")
