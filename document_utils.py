# document_utils.py
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import qn

def setup_document_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for level in [1, 2, 3]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading_style.font.size = Pt(16 - level*2)
        heading_style.font.bold = True

    question_style = doc.styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
    question_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    question_style.font.name = 'Times New Roman'
    question_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_basic_knowledge_section(doc, knowledge_text):
    """
    在试卷前面添加基础知识部分，方便考生温习后续可能涉及的关键知识点
    """
    doc.add_heading("基础知识", level=1)
    doc.add_paragraph(knowledge_text)
    doc.add_page_break()

def add_question_section(doc, question_type, questions, answer_sheet, start_question_count):
    """
    添加题目部分，并更新答案表
    """
    doc.add_heading(f"{question_type}（共{len(questions)}题）", level=2)
    question_count = start_question_count
    for q in questions:
        question_count += 1
        p = doc.add_paragraph(style='Normal')
        p.add_run(f"{question_count}. ").bold = True
        p.add_run(q["question"])
        p.add_run(f"（难度：{q['difficulty']}/5）").italic = True
        answer_sheet.append({
            "number": question_count,
            "type": question_type,
            "question": q["question"],
            "answer": q["short_answer"],
            "analysis": q["detailed_analysis"]
        })
    return question_count

def add_answer_section(doc, answer_sheet):
    """
    添加答案与解析部分
    """
    doc.add_heading("参考答案与解析", level=0)
    for item in answer_sheet:
        doc.add_heading(f"题号{item['number']}（{item['type']}）", level=3)
        doc.add_paragraph(f"题目：{item['question']}")
        answer_para = doc.add_paragraph()
        answer_para.add_run("答案：").bold = True
        answer_para.add_run(item["answer"])
        analysis_para = doc.add_paragraph()
        analysis_para.add_run("解析：").bold = True
        analysis_para.add_run(item["analysis"])
        doc.add_paragraph()

def add_knowledge_summary_section(doc, summary_text):
    """
    将知识点总结文本转换为 docx 格式添加到文档中。
    此函数将 Markdown 格式转换为 docx 格式，去除所有 Markdown 符号，
    并跳过包含答案信息的行（例如以“答案：”开头的行），
    自动根据文本内容设置对应的标题、编号列表或项目符号列表，确保输出为适合 docx 排版的纯文本格式。
    """
    # 添加知识点总结总标题
    doc.add_heading("知识点总结", level=1)
    # 将文本按行拆分
    lines = summary_text.splitlines()
    
    for line in lines:
        # 去除首尾空白
        line = line.strip()
        # 跳过空行或仅由分割符（例如 ---）构成的行
        if not line or re.match(r'^[-_]{3,}$', line):
            continue

        # 先移除 Markdown 特殊符号用于判断
        temp_line = re.sub(r'[#*]', '', line).strip()
        # 如果该行经过清理后以“答案：”开头，则跳过
        if temp_line.startswith("答案："):
            continue

        # 检查是否为 Markdown 标题（如 ### 标记）
        heading_match = re.match(r'^(#{1,6})\s*(.*)', line)
        if heading_match:
            hashes, content = heading_match.groups()
            # 去除内容中的加粗或斜体符号
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            content = re.sub(r'\*(.*?)\*', r'\1', content)
            if content.strip().startswith("答案："):
                continue
            level = min(len(hashes), 4)  # 限制标题级别不超过 4 级
            doc.add_heading(content.strip(), level=level)
            continue

        # 检查是否为数字编号列表（例如 "1. ..."）
        if re.match(r'^\d+\.\s+', line):
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            content = re.sub(r'\*(.*?)\*', r'\1', content)
            if content.strip().startswith("答案："):
                continue
            doc.add_paragraph(content.strip(), style='List Number')
            continue

        # 检查是否为无序列表（例如 "- ..."、"+ ..." 或 "* ..."）
        if re.match(r'^[-+*]\s+', line):
            content = re.sub(r'^[-+*]\s+', '', line)
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            content = re.sub(r'\*(.*?)\*', r'\1', content)
            if content.strip().startswith("答案："):
                continue
            doc.add_paragraph(content.strip(), style='List Bullet')
            continue

        # 其他情况，直接去除加粗、斜体符号后写为普通段落
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        if content.strip().startswith("答案："):
            continue
        doc.add_paragraph(content.strip())

def add_knowledge_summary_section_template(doc, summary_text):
    """
    将固定模板格式的知识点总结解析后添加到 docx 中。
    模板格式要求：每个知识点块之间以 '====' 分隔，每个块内字段为：
    【知识点名称】、【原理】、【实际应用】、【优点】、【缺点】、【注意事项】
    """
    doc.add_heading("知识点总结", level=1)
    blocks = summary_text.split("====")
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        # 将每个块按行拆分，并解析每行的字段
        lines = block.splitlines()
        knowledge_point = {}
        current_field = None
        for line in lines:
            line = line.strip()
            field_match = re.match(r'【(.+?)】：(.*)', line)
            if field_match:
                field, content = field_match.groups()
                knowledge_point[field.strip()] = content.strip()
                current_field = field.strip()
            else:
                if current_field:
                    knowledge_point[current_field] += " " + line
        # 添加到文档中：知识点名称作为二级标题，其余作为段落
        if "知识点名称" in knowledge_point:
            doc.add_heading(knowledge_point["知识点名称"], level=2)
        for key in ["原理", "实际应用", "优点", "缺点", "注意事项"]:
            if key in knowledge_point and knowledge_point[key]:
                para = doc.add_paragraph()
                run = para.add_run(f"{key}：")
                run.bold = True
                para.add_run(knowledge_point[key])
        doc.add_paragraph("")  # 空行分隔
